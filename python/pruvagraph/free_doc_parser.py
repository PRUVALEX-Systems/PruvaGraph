"""
N1 — Free Document Parser

Extracts nodes/edges from PDF, DOCX, Markdown, RST, TXT without any LLM.
Falls back to LLM only for files with zero detectable structure.

Estimated savings: 40–80% of document-type LLM calls eliminated.
"""
from __future__ import annotations

import re
from pathlib import Path


def parse_free(path: Path) -> dict | None:
    """
    Try to parse a document file without LLM.
    Returns dict with nodes/edges, or None if LLM is required.
    """
    ext = path.suffix.lower()

    if ext == ".pdf":
        return _parse_pdf(path)
    if ext == ".docx":
        return _parse_docx(path)
    if ext in (".md", ".markdown", ".mdx"):
        return _parse_markdown(path)
    if ext in (".rst",):
        return _parse_rst(path)
    if ext == ".txt":
        return _parse_plaintext(path)

    return None  # Unknown → LLM


# ── PDF ─────────────────────────────────────────────────────────────────────────

def _parse_pdf(path: Path) -> dict | None:
    try:
        import pypdf  # type: ignore
    except ImportError:
        return None  # pypdf not installed → LLM fallback

    try:
        reader = pypdf.PdfReader(str(path))
    except Exception:
        return None

    # Extract text from first 20 pages
    pages_text = []
    for page in reader.pages[:20]:
        text = page.extract_text() or ""
        if text.strip():
            pages_text.append(text)

    full_text = "\n".join(pages_text)
    if len(full_text.strip()) < 50:
        return None  # Scanned PDF (image-only) → LLM needed

    file_id = str(path)
    first_para = _first_paragraph(full_text)

    nodes = [{
        "id": file_id, "label": path.stem, "type": "doc",
        "file": str(path),
        "summary": first_para or f"PDF document ({len(reader.pages)} pages)",
    }]
    edges = []

    # Extract section headings: uppercase lines or numbered sections
    headings = re.findall(
        r"^(?:\d+\.?\s+)?([A-Z][A-Z\s\-]{4,60})$",
        full_text, re.MULTILINE,
    )
    prev_h1 = file_id
    for heading in headings[:15]:
        heading = heading.strip()
        if len(heading) < 5:
            continue
        h_id = f"{file_id}::{heading}"
        nodes.append({
            "id": h_id, "label": heading, "type": "concept",
            "file": str(path), "summary": f"Section: {heading}",
        })
        edges.append({"source": prev_h1, "target": h_id, "relation": "contains"})

    return {"nodes": nodes, "edges": edges, "source_file": str(path)}


# ── DOCX ─────────────────────────────────────────────────────────────────────────

def _parse_docx(path: Path) -> dict | None:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return None

    try:
        doc = Document(str(path))
    except Exception:
        return None

    file_id = str(path)
    headings = [
        p.text.strip()
        for p in doc.paragraphs
        if p.style.name.startswith("Heading") and p.text.strip()
    ]
    paragraphs = [
        p.text.strip()
        for p in doc.paragraphs
        if not p.style.name.startswith("Heading") and p.text.strip()
    ]

    summary = paragraphs[0][:150] if paragraphs else f"Word document: {path.name}"

    nodes = [{
        "id": file_id, "label": path.stem, "type": "doc",
        "file": str(path), "summary": summary,
    }]
    edges = []

    prev_h1 = file_id
    for h in headings[:12]:
        h_id = f"{file_id}::{h}"
        nodes.append({
            "id": h_id, "label": h, "type": "concept",
            "file": str(path), "summary": f"Section: {h}",
        })
        edges.append({"source": prev_h1, "target": h_id, "relation": "contains"})
        prev_h1 = h_id

    return {"nodes": nodes, "edges": edges, "source_file": str(path)}


# ── Markdown ─────────────────────────────────────────────────────────────────────

def _parse_markdown(path: Path) -> dict:
    content = path.read_text(encoding="utf-8", errors="replace")
    file_id = str(path)

    # Heading hierarchy
    headings: list[tuple[int, str]] = [
        (len(m.group(1)), m.group(2).strip())
        for m in re.finditer(r"^(#{1,4})\s+(.+)$", content, re.MULTILINE)
    ]

    nodes = [{
        "id": file_id, "label": path.stem, "type": "doc",
        "file": str(path),
        "summary": _first_paragraph(content) or f"Markdown: {path.name}",
    }]
    edges = []

    # Stack-based parent tracking
    parent_stack: list[tuple[int, str]] = [(0, file_id)]

    for level, heading in headings[:20]:
        h_id = f"{file_id}::{heading}"
        nodes.append({
            "id": h_id, "label": heading, "type": "concept",
            "file": str(path), "summary": f"Section: {heading}",
        })
        # Pop stack to find correct parent
        while len(parent_stack) > 1 and parent_stack[-1][0] >= level:
            parent_stack.pop()
        parent = parent_stack[-1][1]
        edges.append({"source": parent, "target": h_id, "relation": "contains"})
        parent_stack.append((level, h_id))

    # Code blocks → language usage nodes
    code_langs = set(re.findall(r"^```(\w+)", content, re.MULTILINE))
    for lang in code_langs:
        if lang in ("bash", "sh", "shell", "console", "text", "plain"):
            continue
        lang_id = f"lang:{lang}"
        if not any(n["id"] == lang_id for n in nodes):
            nodes.append({"id": lang_id, "label": lang, "type": "external",
                          "summary": f"Programming language: {lang}"})
        edges.append({"source": file_id, "target": lang_id, "relation": "uses"})

    # Extract links (internal .md links → edges)
    internal_links = re.findall(r"\[([^\]]+)\]\(([^)]+\.md[^)]*)\)", content)
    for link_text, link_href in internal_links[:10]:
        linked_path = path.parent / link_href.split("#")[0]
        edges.append({
            "source": file_id,
            "target": str(linked_path),
            "relation": "references",
        })

    return {"nodes": nodes, "edges": edges, "source_file": str(path)}


# ── RST ─────────────────────────────────────────────────────────────────────────

def _parse_rst(path: Path) -> dict:
    content = path.read_text(encoding="utf-8", errors="replace")
    file_id = str(path)

    # RST headings are underlined with ===, ---, ~~~, etc.
    heading_pattern = re.compile(
        r"^(.+)\n([=\-~`#\"^*+]{3,})\s*$", re.MULTILINE
    )

    nodes = [{
        "id": file_id, "label": path.stem, "type": "doc",
        "file": str(path),
        "summary": _first_paragraph(content) or f"RST document: {path.name}",
    }]
    edges = []

    for m in list(heading_pattern.finditer(content))[:15]:
        heading = m.group(1).strip()
        h_id = f"{file_id}::{heading}"
        nodes.append({
            "id": h_id, "label": heading, "type": "concept",
            "file": str(path), "summary": f"Section: {heading}",
        })
        edges.append({"source": file_id, "target": h_id, "relation": "contains"})

    return {"nodes": nodes, "edges": edges, "source_file": str(path)}


# ── Plain Text ───────────────────────────────────────────────────────────────────

def _parse_plaintext(path: Path) -> dict | None:
    """Only parse TXT if it has enough content worth indexing."""
    content = path.read_text(encoding="utf-8", errors="replace")
    if len(content.strip()) < 100:
        return None  # Too short to be useful

    file_id = str(path)
    summary = _first_paragraph(content)

    return {
        "nodes": [{
            "id": file_id, "label": path.name, "type": "doc",
            "file": str(path),
            "summary": summary or content[:120].replace("\n", " "),
        }],
        "edges": [],
        "source_file": str(path),
    }


# ── Helpers ──────────────────────────────────────────────────────────────────────

def _first_paragraph(text: str) -> str:
    """Extract first meaningful paragraph (non-heading, non-empty)."""
    lines = text.split("\n")
    paras = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if paras:
                break
            continue
        if stripped.startswith("#") or stripped.startswith("="):
            continue
        paras.append(stripped)
        if len(" ".join(paras)) > 150:
            break
    return " ".join(paras)[:200]


DOC_EXTENSIONS = frozenset({
    ".pdf", ".docx", ".md", ".markdown", ".mdx", ".rst", ".txt",
})


def is_parseable_doc(path: Path) -> bool:
    return path.suffix.lower() in DOC_EXTENSIONS
